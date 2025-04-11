import os
import sys
import json
import time
import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException, NoSuchElementException
from docx import Document
from docx.shared import Inches

# Define el directorio del script y del proyecto
SCRIPT_PATH = os.path.dirname(__file__)
PROJECTPATH = os.path.dirname(SCRIPT_PATH)

# Agrega la ruta de 'utils' al sys.path
utils_path = os.path.join(PROJECTPATH, 'utils')
sys.path.append(utils_path)

# Nombre del Test
test_name = "Smoke Test Check SV Trabajador - Mutual.cl"

# Cargar datos desde JSON
data_file = os.path.join(PROJECTPATH, 'utils/input_data', 'data.json')
with open(data_file, 'r') as f:
    data = json.load(f)

urlCertCloud = data['urlCertCloud']

# Credenciales de usuario
#usuario = 'CesarAgurto'
#usuario = 'CarlaMarina'
usuario = 'HernanNava'

rutTrabajadorCert = data['rutTrabajadorCert'][usuario]
claveTrabajadorCert = data['claveTrabajadorCert'][usuario]

# Cargar URLs desde Excel
archivo_excel = os.path.join(PROJECTPATH, 'utils', 'URLs.xlsx')
df = pd.read_excel(archivo_excel, sheet_name="TRAcertCloud")
lista_urls = df['URL'].dropna().tolist()  # Eliminar posibles valores NaN

# Archivos de salida
output_dir = os.path.join(PROJECTPATH, 'utils', 'output_data')
screenshots_dir = os.path.join(output_dir, 'screenshots')
os.makedirs(screenshots_dir, exist_ok=True)  # Crear directorio si no existe

archivo_salida = os.path.join(output_dir, f'Evidencias_Check_SVT_{usuario}.docx')

# Crear documento Word
doc = Document()
doc.add_heading(f'Resultados de ejecución - {test_name}', level=1)

def take_screenshot(driver, file_path):
    """ Captura la pantalla y la guarda en la ruta especificada. """
    driver.save_screenshot(file_path)

def check_banner(driver, url, idx):
    """ Verifica si el banner está visible, guarda su texto y captura su imagen. """
    banner_id = "titulo-header-banner"
    try:
        banner = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.ID, banner_id)))
        if banner.is_displayed():
            banner_text = banner.text.strip() if banner.text.strip() else "⚠️ Sin texto visible"
            doc.add_paragraph(f"✅ Banner visible en {url}.")
            doc.add_paragraph(f"📌 Texto del banner: {banner_text}")
            
            # Captura del banner
            banner_screenshot = os.path.join(screenshots_dir, f"screenshot_{idx}_banner.png")
            take_screenshot(driver, banner_screenshot)
            doc.add_paragraph("📷 Captura del banner:")
            doc.add_picture(banner_screenshot, width=Inches(5))
        else:
            doc.add_paragraph(f"⚠️ Banner encontrado en {url}, pero no está visible.")
    except (TimeoutException, NoSuchElementException):
        doc.add_paragraph(f"❌ Banner no encontrado en {url}.")

def capture_full_page(driver, url, idx):
    """ Captura toda la página dividiéndola en varias secciones si es necesario. """
    try:
        # Obtener dimensiones de la página
        page_height = driver.execute_script("return document.body.scrollHeight")
        viewport_height = driver.execute_script("return window.innerHeight")

        # Calcular el número de capturas necesarias
        num_screenshots = max(1, (page_height // viewport_height) + 1)

        print(f"📸 Se tomarán {num_screenshots} capturas para {url}")

        for i in range(num_screenshots):
            # Calcular la posición del scroll
            scroll_position = i * viewport_height
            driver.execute_script(f"window.scrollTo(0, {scroll_position})")
            time.sleep(1)  # Breve espera para permitir el renderizado

            # Capturar y guardar la imagen
            screenshot_path = os.path.join(screenshots_dir, f"screenshot_{idx}_part_{i+1}.png")
            take_screenshot(driver, screenshot_path)

            # Agregar la captura al documento
            doc.add_paragraph(f"📷 Captura {i+1} de {num_screenshots} de {url}:")
            doc.add_picture(screenshot_path, width=Inches(5))

    except Exception as e:
        print(f"⚠️ Error capturando la página {url}: {e}")
        doc.add_paragraph(f"⚠️ No se pudo capturar correctamente la página {url}")

if __name__ == "__main__":
    print(f"Iniciando el test: {test_name}")

    # Importar y llamar a gotoMutual
    from gotoMutual import gotoMutual
    driver = gotoMutual(urlCertCloud)

    # Importar y llamar a loginSVTMutual
    from loginSVTMutual import loginSVTMutual
    driver = loginSVTMutual(driver, rutTrabajadorCert, claveTrabajadorCert)

    # Guardar la pestaña principal
    main_window = driver.current_window_handle

    # Recorrer las URLs
    for idx, url in enumerate(lista_urls, start=1):
        try:
            print(f"🟢 Visitando: {url}")
            driver.execute_script("window.open('', '_blank');")
            driver.switch_to.window(driver.window_handles[-1])
            driver.get(url)
            titulo = driver.title

            # Espera a que la página cargue
            try:
                WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
                doc.add_paragraph(f"✅ {idx}. {url} - {titulo} cargó correctamente.")
            except TimeoutException:
                doc.add_paragraph(f"⚠️ {idx}. {url} - {titulo} tardó demasiado en cargar.")

            # Capturar toda la página
            capture_full_page(driver, url, idx)

            # Cerrar pestaña actual y volver a la principal
            driver.close()
            driver.switch_to.window(main_window)

        except WebDriverException as e:
            print(f"🔴 Error con {url}: {e}")
            doc.add_paragraph(f"❌ Error con {idx}. {url} - {titulo}: {e}")

    # Guardar documento de evidencias
    doc.save(archivo_salida)
    print(f"\n📄 Resultados guardados en '{archivo_salida}'")

    # Cerrar navegador
    time.sleep(3)
    driver.quit()
    print(f"Test finalizado: {test_name}")
# Fin del script
