import os
import sys
import json
import time
import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException, NoSuchElementException
from docx import Document
from docx.shared import Inches, RGBColor
from datetime import datetime

# Rutas y configuración
SCRIPT_PATH = os.path.dirname(__file__)
PROJECTPATH = os.path.dirname(SCRIPT_PATH)
utils_path = os.path.join(PROJECTPATH, 'utils')
sys.path.append(utils_path)

test_name = "Revisar Site Rediseño Mutual"
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

# Cargar datos
with open(os.path.join(PROJECTPATH, 'utils/input_data', 'data.json'), 'r') as f:
    data = json.load(f)
urlCertRediseno = data['urlCertRediseno']
urlProdRediseno = data['urlProdRediseno']

# Leer URLs desde Excel
archivo_excel = os.path.join(PROJECTPATH, 'utils', 'URLs.xlsx')
pestaña = "RedisenoPROD"  # O "RedisenoCERT"
df = pd.read_excel(archivo_excel, sheet_name=pestaña)
lista_urls = df['URL'].dropna().tolist()

# Salidas
output_dir = os.path.join(PROJECTPATH, 'utils', 'output_data')
screenshots_dir = os.path.join(output_dir, f'screenshots_{timestamp}')
os.makedirs(screenshots_dir, exist_ok=True)
archivo_docx = os.path.join(output_dir, f'Evidencias_Ejecucion_{pestaña}_{timestamp}.docx')
archivo_excel_resumen = os.path.join(output_dir, f'Resumen_resultados_{pestaña}_{timestamp}.xlsx')

# Documento y resumen
doc = Document()
doc.add_heading(f'Resultados de ejecución - {test_name}', level=1)
resumen_resultados = []
referencias_detectadas = []

# Palabras clave sospechosas
palabras_prohibidas = [
     "https://qa.", "https://cert.", "https://172.177.177.139/"
]

def take_screenshot(driver, file_path):
    driver.save_screenshot(file_path)

def check_banner(driver, url, idx):
    try:
        banner = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.ID, "titulo-header-banner")))
        banner_text = banner.text.strip() if banner.text.strip() else ""

        if banner_text:
            doc.add_paragraph(f"✅ Banner: {banner_text}")
            return "✅ OK", "Banner visible"
        else:
            run = doc.add_paragraph().add_run(f"⚠️ Banner sin texto en {url}.")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 140, 0)
            return "⚠️ Advertencia", "Banner sin texto"
    except (TimeoutException, NoSuchElementException):
        run = doc.add_paragraph().add_run(f"❌ Banner no encontrado en {url}.")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        return "❌ Error", "Banner no encontrado"

def check_error_message(driver, url, idx):
    try:
        error_block = WebDriverWait(driver, 3).until(EC.presence_of_element_located((By.CLASS_NAME, "lrpErrorContent")))
        h1 = error_block.find_element(By.TAG_NAME, "h1").text.strip()
        p = error_block.find_element(By.TAG_NAME, "p").text.strip()

        doc.add_paragraph().add_run(f"❌ Error en {url}:").bold = True
        doc.add_paragraph(f"🔹 Título: {h1}")
        doc.add_paragraph(f"🔸 Descripción: {p}")
        take_screenshot(driver, os.path.join(screenshots_dir, f"screenshot_{idx}_error.png"))
        return "❌ Error", f"{h1} - {p[:40]}..."
    except TimeoutException:
        return "✅ OK", "Sin error visible"
    except NoSuchElementException:
        return "⚠️ Advertencia", "Contenedor de error sin contenido"

def buscar_referencias_html(driver, palabras_clave, url, idx):
    html = driver.page_source.lower()
    encontrados = [p for p in palabras_clave if p in html]
    if encontrados:
        doc.add_paragraph().add_run(f"⚠️ Referencias sospechosas en {url}:").bold = True
        doc.add_paragraph(f"🔍 Coincidencias: {', '.join(encontrados)}")
        take_screenshot(driver, os.path.join(screenshots_dir, f"screenshot_{idx}_referencias.png"))
        detalle = f"Referencias: {', '.join(encontrados)}"
        referencias_detectadas.append({"idx": idx, "url": url, "coincidencias": detalle})
        return "⚠️ Advertencia", detalle
    return "✅ OK", "Sin referencias sospechosas"

def verificar_contenedor_contenido(driver, by, value, nombre_visible, idx):
    try:
        contenedor = driver.find_element(by, value)
        texto = contenedor.text.strip()
        hijos = contenedor.find_elements(By.XPATH, "./*")
        if texto:
            doc.add_paragraph(f"✅ {nombre_visible} con texto: {texto[:80]}...")
        elif hijos:
            doc.add_paragraph(f"⚠️ {nombre_visible} sin texto, pero con {len(hijos)} hijos.")
        else:
            run = doc.add_paragraph().add_run(f"⚠️ {nombre_visible} está vacío.")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 140, 0)
            return "⚠️ Advertencia", f"{nombre_visible} vacío"
        return "✅ OK", f"{nombre_visible} con contenido"
    except NoSuchElementException:
        run = doc.add_paragraph().add_run(f"❌ {nombre_visible} no encontrado.")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        return "❌ Error", f"{nombre_visible} no encontrado"

def capture_full_page(driver, url, idx):
    height = driver.execute_script("return document.body.scrollHeight")
    view = driver.execute_script("return window.innerHeight")
    total = max(1, (height + view - 1) // view)
    for i in range(total):
        driver.execute_script(f"window.scrollTo(0, {i * view})")
        WebDriverWait(driver, 2).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        take_screenshot(driver, os.path.join(screenshots_dir, f"screenshot_{idx}_part_{i+1}.png"))

# Ejecución principal
if __name__ == "__main__":
    from gotoMutual import gotoMutual
    driver = gotoMutual(urlProdRediseno)

    for idx, url in enumerate(lista_urls, start=1):
        try:
            start = time.time()
            driver.get(url)
            load_time = round(time.time() - start, 2)
            doc.add_paragraph(f"✅ {idx}. {url} cargó en {load_time} seg.")

            estado_banner, detalle_banner = check_banner(driver, url, idx)
            estado_error, detalle_error = check_error_message(driver, url, idx)
            estado_ref, detalle_ref = buscar_referencias_html(driver, palabras_prohibidas, url, idx)
            estado_contenedor, detalle_contenedor = verificar_contenedor_contenido(driver, By.NAME, "ibmMainContainer", "Contenedor Principal", idx)
            capture_full_page(driver, url, idx)

            # Prioridad de estados: ❌ > ⚠️ > ✅
            if estado_error.startswith("❌"):
                estado, detalle = estado_error, detalle_error
            elif estado_ref.startswith("⚠️"):
                estado, detalle = estado_ref, detalle_ref
            elif estado_contenedor.startswith("❌") or estado_contenedor.startswith("⚠️"):
                estado, detalle = estado_contenedor, detalle_contenedor
            elif estado_banner.startswith("❌"):
                estado, detalle = estado_banner, detalle_banner
            elif estado_banner.startswith("⚠️") or estado_error.startswith("⚠️") or estado_contenedor.startswith("⚠️"):
                estado = "⚠️ Advertencia"
                detalle = detalle_error if estado_error.startswith("⚠️") else detalle_banner
            else:
                estado, detalle = "✅ OK", "Carga exitosa"

            resumen_resultados.append([idx, url, estado, load_time, detalle])

        except WebDriverException as e:
            doc.add_paragraph().add_run(f"❌ Error con {idx}. {url}: {e}").bold = True
            resumen_resultados.append([idx, url, "❌ Error", 0, str(e)])

    # Agregar resumen al Word
    doc.add_page_break()
    doc.add_heading("Resumen de resultados", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    for i, h in enumerate(["#", "URL", "Estado", "Tiempo (s)", "Detalles"]):
        table.rows[0].cells[i].text = h
    for fila in resumen_resultados:
        row = table.add_row().cells
        for i, val in enumerate(fila):
            row[i].text = str(val)

    if referencias_detectadas:
        doc.add_page_break()
        doc.add_heading("URLs con referencias sospechosas", level=2)
        ref_table = doc.add_table(rows=1, cols=3)
        ref_table.style = 'Light Grid Accent 2'
        ref_table.rows[0].cells[0].text = "#"
        ref_table.rows[0].cells[1].text = "URL"
        ref_table.rows[0].cells[2].text = "Coincidencias"
        for ref in referencias_detectadas:
            row = ref_table.add_row().cells
            row[0].text = str(ref['idx'])
            row[1].text = ref['url']
            row[2].text = ref['coincidencias']

    doc.save(archivo_docx)

    # Guardar Excel
    resumen_df = pd.DataFrame(resumen_resultados, columns=["#", "URL", "Estado", "Tiempo (s)", "Detalles"])
    referencias_df = pd.DataFrame(referencias_detectadas)
    with pd.ExcelWriter(archivo_excel_resumen, engine='xlsxwriter') as writer:
        resumen_df.to_excel(writer, sheet_name='Resumen', index=False)
        if not referencias_df.empty:
            referencias_df.to_excel(writer, sheet_name='Referencias sospechosas', index=False)

    driver.quit()
    print(f"\n📄 Evidencias Word: {archivo_docx}")
    print(f"📊 Excel con dos hojas: {archivo_excel_resumen}")