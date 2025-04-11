import os
import sys
import json
import time
import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException, NoSuchElementException
from docx import Document
from docx.shared import Inches
from datetime import datetime

# Definir rutas
SCRIPT_PATH = os.path.dirname(__file__)
PROJECTPATH = os.path.dirname(SCRIPT_PATH)

# Agregar la ruta de 'utils'
utils_path = os.path.join(PROJECTPATH, 'utils')
sys.path.append(utils_path)

# Configuración de ejecución
test_name = "Revisar Site Rediseño - Mutual.cl"
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

# Cargar datos
data_file = os.path.join(PROJECTPATH, 'utils/input_data', 'data.json')
with open(data_file, 'r') as f:
    data = json.load(f)
urlCertRediseno = data['urlCertRediseno']

# Cargar URLs desde Excel
archivo_excel = os.path.join(PROJECTPATH, 'utils', 'URLs.xlsx')
df = pd.read_excel(archivo_excel, sheet_name="Rediseno")
lista_urls = df['URL'].dropna().tolist()

# Crear carpetas de salida
output_dir = os.path.join(PROJECTPATH, 'utils', 'output_data')
screenshots_dir = os.path.join(output_dir, f'screenshots_{timestamp}')
os.makedirs(screenshots_dir, exist_ok=True)

archivo_salida = os.path.join(output_dir, f'Evidencias_Ejecucion_Rediseno_{timestamp}.docx')

# Documento Word
doc = Document()
doc.add_heading(f'Resultados de ejecución - {test_name}', level=1)

def take_screenshot(driver, file_path):
    driver.save_screenshot(file_path)

def check_banner(driver, url, idx):
    try:
        banner = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.ID, "titulo-header-banner")))
        if banner.is_displayed():
            banner_text = banner.text.strip() if banner.text.strip() else "⚠️ Sin texto visible"
            doc.add_paragraph(f"✅ Banner: {banner_text}")
            banner_screenshot = os.path.join(screenshots_dir, f"screenshot_{idx}_banner.png")
            take_screenshot(driver, banner_screenshot)
            doc.add_picture(banner_screenshot, width=Inches(5))
    except (TimeoutException, NoSuchElementException):
        doc.add_paragraph(f"❌ Banner no encontrado en {url}.")

def capture_full_page(driver, url, idx):
    page_height = driver.execute_script("return document.body.scrollHeight")
    viewport_height = driver.execute_script("return window.innerHeight")
    num_screenshots = max(1, (page_height + viewport_height - 1) // viewport_height)

    for i in range(num_screenshots):
        driver.execute_script(f"window.scrollTo(0, {i * viewport_height})")
        WebDriverWait(driver, 2).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        screenshot_path = os.path.join(screenshots_dir, f"screenshot_{idx}_part_{i+1}.png")
        take_screenshot(driver, screenshot_path)
        doc.add_picture(screenshot_path, width=Inches(5))

if __name__ == "__main__":
    from gotoMutual import gotoMutual
    driver = gotoMutual(urlCertRediseno)
    main_window = driver.current_window_handle

    for idx, url in enumerate(lista_urls, start=1):
        try:
            start_time = time.time()
            driver.get(url)
            load_time = round(time.time() - start_time, 2)
            doc.add_paragraph(f"✅ {idx}. {url} cargó en {load_time} seg.")

            check_banner(driver, url, idx)
            capture_full_page(driver, url, idx)

        except WebDriverException as e:
            doc.add_paragraph(f"❌ Error con {idx}. {url}: {e}")

    doc.save(archivo_salida)
    driver.quit()
    print(f"\n📄 Evidencias guardadas en '{archivo_salida}'")

