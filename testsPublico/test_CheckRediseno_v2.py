import os
import sys
import json
import time
import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException, NoSuchElementException
from docx import Document
from docx.shared import Inches, RGBColor
from datetime import datetime

# Rutas
SCRIPT_PATH = os.path.dirname(__file__)
PROJECTPATH = os.path.dirname(SCRIPT_PATH)
utils_path = os.path.join(PROJECTPATH, 'utils')
sys.path.append(utils_path)

# Configuración
test_name = "Revisar Site Rediseño Mutual"
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

# Cargar datos
with open(os.path.join(PROJECTPATH, 'utils/input_data', 'data.json'), 'r') as f:
    data = json.load(f)
urlCertRediseno = data['urlCertRediseno']
urlProdRediseno = data['urlProdRediseno']

# Leer Excel
archivo_excel = os.path.join(PROJECTPATH, 'utils', 'URLs.xlsx')
# pestaña = "RedisenoCERT"
pestaña = "RedisenoPROD"
df = pd.read_excel(archivo_excel, sheet_name=pestaña)
lista_urls = df['URL'].dropna().tolist()

# Salidas
output_dir = os.path.join(PROJECTPATH, 'utils', 'output_data')
screenshots_dir = os.path.join(output_dir, f'screenshots_{timestamp}')
os.makedirs(screenshots_dir, exist_ok=True)
archivo_docx = os.path.join(output_dir, f'Evidencias_Ejecucion_{pestaña}_{timestamp}.docx')
archivo_excel_resumen = os.path.join(output_dir, f'Resumen_resultados_{pestaña}_{timestamp}.xlsx')

doc = Document()
doc.add_heading(f'Resultados de ejecución - {test_name}', level=1)
resumen_resultados = []

def take_screenshot(driver, file_path):
    driver.save_screenshot(file_path)

def check_banner(driver, url, idx, retry=True):
    try:
        banner = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.ID, "titulo-header-banner")))
        banner_text = banner.text.strip() if banner.text.strip() else ""

        if not banner_text and retry:
            driver.refresh()
            time.sleep(3)
            return check_banner(driver, url, idx, retry=False)

        if banner_text:
            doc.add_paragraph(f"✅ Banner: {banner_text}")
            return "✅ OK", "Banner visible"
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"⚠️ Banner encontrado pero sin texto visible en {url}.")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 140, 0)
            return "⚠️ Advertencia", "Banner sin texto"

    except (TimeoutException, NoSuchElementException):
        p = doc.add_paragraph()
        run = p.add_run(f"❌ Banner no encontrado en {url}.")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        return "❌ Error", "Banner no encontrado"

def check_error_message(driver, url, idx):
    try:
        error_block = WebDriverWait(driver, 3).until(
            EC.presence_of_element_located((By.CLASS_NAME, "lrpErrorContent"))
        )

        h1_element = error_block.find_element(By.TAG_NAME, "h1")
        p_element = error_block.find_element(By.TAG_NAME, "p")

        h1_text = h1_element.text.strip() if h1_element else "Sin título"
        p_text = p_element.text.strip() if p_element else "Sin descripción"

        doc.add_paragraph().add_run(f"❌ Error encontrado en {url}:\n").bold = True
        doc.add_paragraph().add_run(f"🔹 Título: {h1_text}").font.color.rgb = RGBColor(255, 0, 0)
        doc.add_paragraph().add_run(f"🔸 Descripción: {p_text}").font.color.rgb = RGBColor(255, 0, 0)

        error_screenshot = os.path.join(screenshots_dir, f"screenshot_{idx}_error.png")
        take_screenshot(driver, error_screenshot)
        doc.add_picture(error_screenshot, width=Inches(5))

        return "❌ Error", f"{h1_text} - {p_text[:40]}..."

    except TimeoutException:
        return "✅ OK", "Sin mensaje de error"
    except NoSuchElementException:
        return "⚠️ Advertencia", "Contenedor de error sin texto"

def capture_full_page(driver, url, idx):
    page_height = driver.execute_script("return document.body.scrollHeight")
    viewport_height = driver.execute_script("return window.innerHeight")
    num_screenshots = max(1, (page_height + viewport_height - 1) // viewport_height)

    for i in range(num_screenshots):
        driver.execute_script(f"window.scrollTo(0, {i * viewport_height})")
        WebDriverWait(driver, 2).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        screenshot_path = os.path.join(screenshots_dir, f"screenshot_{idx}_part_{i+1}.png")
        take_screenshot(driver, screenshot_path)
        doc.add_picture(screenshot_path, width=Inches(5))

if __name__ == "__main__":
    from gotoMutual import gotoMutual
    # driver = gotoMutual(urlCertRediseno)
    driver = gotoMutual(urlProdRediseno)

    for idx, url in enumerate(lista_urls, start=1):
        try:
            start_time = time.time()
            driver.get(url)
            load_time = round(time.time() - start_time, 2)
            doc.add_paragraph(f"✅ {idx}. {url} cargó en {load_time} seg.")

            estado_banner, detalle_banner = check_banner(driver, url, idx)
            estado_error, detalle_error = check_error_message(driver, url, idx)
            capture_full_page(driver, url, idx)

            # Prioridad del estado: ❌ > ⚠️ > ✅
            if estado_error.startswith("❌"):
                estado, detalle = estado_error, detalle_error
            elif estado_banner.startswith("❌"):
                estado, detalle = estado_banner, detalle_banner
            elif estado_error.startswith("⚠️") or estado_banner.startswith("⚠️"):
                estado = "⚠️ Advertencia"
                detalle = detalle_error if estado_error.startswith("⚠️") else detalle_banner
            else:
                estado, detalle = "✅ OK", "Carga exitosa"

            resumen_resultados.append([idx, url, estado, load_time, detalle])

        except WebDriverException as e:
            run = doc.add_paragraph().add_run(f"❌ Error con {idx}. {url}: {e}")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            resumen_resultados.append([idx, url, "❌ Error", 0, str(e)])

    # Tabla en Word
    doc.add_page_break()
    doc.add_heading("Resumen de resultados", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    for i, h in enumerate(["#", "URL", "Estado", "Tiempo (s)", "Detalles"]):
        table.rows[0].cells[i].text = h
    for fila in resumen_resultados:
        row_cells = table.add_row().cells
        for i, val in enumerate(fila):
            row_cells[i].text = str(val)

    # Guardar documentos
    doc.save(archivo_docx)
    pd.DataFrame(resumen_resultados, columns=["#", "URL", "Estado", "Tiempo (s)", "Detalles"]).to_excel(archivo_excel_resumen, index=False)

    driver.quit()
    print(f"\n📄 Evidencias Word: {archivo_docx}")
    print(f"📊 Resumen en Excel: {archivo_excel_resumen}")
